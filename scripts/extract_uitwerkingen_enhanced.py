"""
Enhanced DOCX Extractor - Preserves Inline Images and Formatting

This script extracts content from the Uitwerkingen DOCX with:
- Full text content preservation
- Inline image position tracking
- Text formatting (bold, italic)
- Paragraph structure
- Formula/image inline markers

Output: analysis/uitwerkingen-enhanced.yaml (with inline image positions)
"""

import sys
import re
from pathlib import Path
from typing import List, Dict, Any, Optional
from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.text.run import Run
import yaml

# Force UTF-8 encoding
if sys.platform == 'win32':
    sys.stdout.reconfigure(encoding='utf-8')
    sys.stderr.reconfigure(encoding='utf-8')

# Paths
SCRIPT_DIR = Path(__file__).parent
PROJECT_ROOT = SCRIPT_DIR.parent
INPUT_DOCX = PROJECT_ROOT / "source-materials" / "Uitwerkingen Opgaven Energietechniek V0.8.docx"
OUTPUT_YAML = PROJECT_ROOT / "analysis" / "uitwerkingen-enhanced.yaml"
OUTPUT_MD = PROJECT_ROOT / "analysis" / "uitwerkingen-enhanced.md"
IMAGES_DIR = PROJECT_ROOT / "images" / "uitwerkingen"

# Make sure images dir exists
IMAGES_DIR.mkdir(parents=True, exist_ok=True)

# Global image counter
image_counter = 0


def extract_images_from_paragraph(paragraph: Paragraph, para_index: int) -> List[Dict[str, Any]]:
    """
    Extract ALL images from a paragraph including formulas.
    Uses comprehensive extraction similar to old extractor.
    """
    global image_counter
    images = []

    # Extract from ALL relationship types (like old extractor did)
    for run in paragraph.runs:
        # Check if run contains images (comprehensive check)
        if 'graphic' in run._element.xml or 'pict' in run._element.xml or 'drawing' in run._element.xml:
            # Get ALL image relationships
            for rel in paragraph.part.rels.values():
                if "image" in rel.target_ref:
                    try:
                        image_data = rel.target_part.blob

                        # Determine image extension
                        content_type = rel.target_part.content_type
                        if 'png' in content_type:
                            ext = 'png'
                        elif 'jpeg' in content_type or 'jpg' in content_type:
                            ext = 'jpg'
                        elif 'gif' in content_type:
                            ext = 'gif'
                        elif 'emf' in content_type or 'wmf' in content_type:
                            ext = 'emf'  # Windows metafile (formulas!)
                        else:
                            ext = 'png'

                        # Save image
                        image_counter += 1
                        filename = f"opgave_{image_counter:04d}.{ext}"
                        image_path = IMAGES_DIR / filename

                        with open(image_path, 'wb') as f:
                            f.write(image_data)

                        # Track position (approximate)
                        char_pos = sum(len(r.text) for r in paragraph.runs[:paragraph.runs.index(run)])

                        images.append({
                            'position': char_pos,
                            'filename': filename,
                            'type': 'inline',
                            'content_type': content_type
                        })

                    except Exception as e:
                        pass  # Silent fail

    return images


def extract_paragraph_with_formatting(paragraph: Paragraph, para_index: int) -> Dict[str, Any]:
    """
    Extract paragraph content with formatting and inline images.

    Returns dict with:
    - text: plain text
    - runs: list of formatted text runs
    - images: list of inline images with positions
    - style: paragraph style
    """
    runs_data = []
    full_text = ""

    for run in paragraph.runs:
        run_data = {
            'text': run.text,
            'bold': run.bold,
            'italic': run.italic,
            'underline': run.underline,
        }
        runs_data.append(run_data)
        full_text += run.text

    # Extract images with positions
    images = extract_images_from_paragraph(paragraph, para_index)

    return {
        'text': full_text.strip(),
        'runs': runs_data,
        'images': images,
        'style': paragraph.style.name if paragraph.style else 'Normal'
    }


def detect_section_type(text: str) -> Optional[str]:
    """Detect if this is an Opgave or Uitwerking header."""
    text_lower = text.lower().strip()

    # Opgave patterns
    if re.match(r'^opgave\s+', text_lower):
        return 'opgave_header'

    # Uitwerking patterns
    if re.match(r'^uitwerking\s+', text_lower):
        return 'uitwerking_header'

    # Chapter patterns
    if re.match(r'^h\s+\d+\.\d+', text_lower):
        return 'chapter_header'

    return None


def extract_enhanced_docx(docx_path: Path) -> List[Dict[str, Any]]:
    """
    Extract DOCX with full content preservation.

    Returns list of structured sections with inline images.
    """
    print(f"\n{'='*70}")
    print("ENHANCED DOCX EXTRACTION")
    print(f"{'='*70}\n")

    print(f"Reading: {docx_path.name}")
    doc = Document(docx_path)

    print(f"✓ Document loaded: {len(doc.paragraphs)} paragraphs")

    # Structure to hold all content
    sections = []
    current_section = None
    current_chapter = None

    for para_idx, paragraph in enumerate(doc.paragraphs):
        if para_idx % 100 == 0:
            print(f"  Processing paragraph {para_idx}/{len(doc.paragraphs)}...")

        # Extract paragraph with formatting
        para_data = extract_paragraph_with_formatting(paragraph, para_idx)

        # Skip empty paragraphs
        if not para_data['text'] and not para_data['images']:
            continue

        # Detect section type
        section_type = detect_section_type(para_data['text'])

        # Handle chapter headers
        if section_type == 'chapter_header':
            match = re.match(r'H\s+(\d+)\.(\d+)\s+blz\.?\s+(\d+)', para_data['text'], re.IGNORECASE)
            if match:
                current_chapter = {
                    'chapter': int(match.group(1)),
                    'section': int(match.group(2)),
                    'page': int(match.group(3)),
                    'title': para_data['text']
                }
                print(f"  📖 Chapter: H{current_chapter['chapter']}.{current_chapter['section']}")
            continue

        # Handle opgave headers
        if section_type == 'opgave_header':
            # Save previous section if exists
            if current_section:
                sections.append(current_section)

            # Start new section
            current_section = {
                'type': 'opgave',
                'header': para_data['text'],
                'chapter': current_chapter,
                'content': [],
                'opgave_paragraphs': [],
                'uitwerking_paragraphs': [],
                'current_mode': 'opgave'
            }
            continue

        # Handle uitwerking headers
        if section_type == 'uitwerking_header':
            if current_section:
                current_section['current_mode'] = 'uitwerking'
            continue

        # Add content to current section
        if current_section:
            para_data['paragraph_index'] = para_idx

            if current_section['current_mode'] == 'opgave':
                current_section['opgave_paragraphs'].append(para_data)
            else:
                current_section['uitwerking_paragraphs'].append(para_data)

            current_section['content'].append(para_data)

    # Save last section
    if current_section:
        sections.append(current_section)

    print(f"\n✓ Extracted {len(sections)} sections")
    print(f"✓ Extracted {image_counter} images")

    return sections


def sections_to_structured_problems(sections: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """Convert sections to structured problem format."""

    problems = []

    for section in sections:
        if section['type'] != 'opgave':
            continue

        # Extract opgave ID from header
        header_text = section['header']
        match = re.match(r'Opgave\s+([0-9]+[a-z]?)', header_text, re.IGNORECASE)
        opgave_id = match.group(1) if match else "unknown"

        # Build problem structure
        problem = {
            'id': opgave_id,
            'header': header_text,
            'chapter_ref': section['chapter'],
            'opgave_paragraphs': section['opgave_paragraphs'],
            'uitwerking_paragraphs': section['uitwerking_paragraphs'],
            'total_images': sum(
                len(p['images']) for p in section['content']
            ),
            'total_paragraphs': len(section['content'])
        }

        problems.append(problem)

    return problems


def generate_enhanced_markdown(problems: List[Dict[str, Any]], output_file: Path):
    """Generate markdown with inline image markers."""

    lines = []
    lines.append("# Uitwerkingen Opgaven - Enhanced Extraction\n")
    lines.append(f"**Total Problems:** {len(problems)}\n")
    lines.append(f"**Extraction Date:** 2025-11-08\n\n")
    lines.append("---\n\n")

    for problem in problems:
        lines.append(f"## Opgave {problem['id']}\n\n")

        if problem['chapter_ref']:
            ch = problem['chapter_ref']
            lines.append(f"**Hoofdstuk:** H {ch['chapter']}.{ch['section']} (blz. {ch['page']})\n\n")

        # Opgave content
        lines.append("### Vraagstelling\n\n")
        for para in problem['opgave_paragraphs']:
            # Insert text with inline image markers
            text = para['text']

            # Insert image markers at positions
            if para['images']:
                # Sort images by position (descending to insert from end)
                sorted_images = sorted(para['images'], key=lambda x: x['position'], reverse=True)
                for img in sorted_images:
                    marker = f" ![inline]({img['filename']}) "
                    text = text[:img['position']] + marker + text[img['position']:]

            lines.append(f"{text}\n\n")

        # Uitwerking content
        if problem['uitwerking_paragraphs']:
            lines.append("### Uitwerking\n\n")
            for para in problem['uitwerking_paragraphs']:
                text = para['text']

                # Insert inline images
                if para['images']:
                    sorted_images = sorted(para['images'], key=lambda x: x['position'], reverse=True)
                    for img in sorted_images:
                        marker = f" ![inline]({img['filename']}) "
                        text = text[:img['position']] + marker + text[img['position']:]

                lines.append(f"{text}\n\n")

        lines.append("---\n\n")

    with open(output_file, 'w', encoding='utf-8') as f:
        f.writelines(lines)

    print(f"✓ Markdown written to: {output_file.relative_to(PROJECT_ROOT)}")


def main():
    """Main execution."""

    if not INPUT_DOCX.exists():
        print(f"❌ Error: DOCX file not found: {INPUT_DOCX}")
        return

    # Extract with enhanced tracking
    sections = extract_enhanced_docx(INPUT_DOCX)

    # Convert to structured problems
    problems = sections_to_structured_problems(sections)

    print(f"\n✓ Structured {len(problems)} problems")

    # Save to YAML
    print(f"\nSaving structured data...")
    with open(OUTPUT_YAML, 'w', encoding='utf-8') as f:
        yaml.dump(problems, f, allow_unicode=True, default_flow_style=False, sort_keys=False)

    print(f"✓ YAML written to: {OUTPUT_YAML.relative_to(PROJECT_ROOT)}")

    # Generate markdown preview
    generate_enhanced_markdown(problems, OUTPUT_MD)

    print(f"\n{'='*70}")
    print("EXTRACTION COMPLETE")
    print(f"{'='*70}")
    print(f"\nOutput files:")
    print(f"  - {OUTPUT_YAML.relative_to(PROJECT_ROOT)}")
    print(f"  - {OUTPUT_MD.relative_to(PROJECT_ROOT)}")
    print(f"  - {IMAGES_DIR.relative_to(PROJECT_ROOT)}/ ({image_counter} images)")

    print(f"\nNext steps:")
    print(f"  1. Review: {OUTPUT_MD.relative_to(PROJECT_ROOT)}")
    print(f"  2. Validate inline images are correctly positioned")
    print(f"  3. Generate HTML study guide with inline images")


if __name__ == "__main__":
    main()
